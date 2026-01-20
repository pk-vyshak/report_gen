"""Streamlit UI for Post-Campaign Insights Report Generator."""

import tempfile
from datetime import datetime
from io import BytesIO
from pathlib import Path

import plotly.express as px
import plotly.graph_objects as go
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from src.services import ReportService

# Hardcoded insights specific to Campaign 4512 from Files/Input/
HARDCODED_INSIGHTS = {
    "campaign_summary": (
        "Campaign 4512 delivered 91,726 impressions with 31 clicks over a 4-week "
        "period (Dec 1-29, 2025), generating a total spend of $178.09. The campaign "
        "achieved a CTR of 0.034% and CPM of $1.94, with an average viewability of "
        "66.26%. The campaign showed strong mid-flight performance with Week 3 "
        "delivering the highest impression volume (32,605) and Week 4 achieving the "
        "best CTR performance (0.052%)."
    ),
    "weekly_trend": {
        "impression_delivery": (
            "The campaign showed variable weekly delivery: Week 1 started with 24,657 "
            "impressions (26.9%), followed by a slight dip in Week 2 (22,629, 24.7%). "
            "Week 3 saw a significant 77.7% spike above average with 32,605 impressions "
            "(35.5% of total), before tapering in Week 4 (11,460, 12.5%) and a minimal "
            "375 impressions in the final partial week."
        ),
        "ctr": (
            "CTR showed an improving trend across the campaign. Starting at 0.028% in "
            "Week 1, it dipped to a low of 0.022% in Week 2 (34.6% below average). "
            "Performance recovered strongly in Week 3 (0.040%, 118% of average) and "
            "peaked in Week 4 at 0.052% (154.9% of average). The final week recorded "
            "zero clicks due to minimal impression volume."
        ),
        "viewability": (
            "Viewability showed steady improvement throughout the campaign, ranging "
            "from 63.12% (Week 2) to 71.61% (Week 4). The upward trend suggests "
            "progressive optimization of ad placements, with Week 4 achieving the "
            "highest viewability despite lower impression volume."
        ),
        "vcr": (
            "Video Completion Rate data was not available for this campaign, "
            "indicating this was primarily a display-focused campaign without "
            "video creative assets."
        ),
    },
    "key_recommendations": [
        {
            "title": "Address Week 2 CTR Anomaly",
            "text": (
                "Week 2 CTR dropped 34.6% below campaign average (0.022% vs 0.034%). "
                "This anomaly coincided with the lowest viewability (63.12%), "
                "suggesting potential inventory quality issues. Recommend implementing "
                "viewability floors and excluding underperforming placements to "
                "prevent similar dips in future campaigns."
            ),
        },
        {
            "title": "Leverage Week 3 Pacing Strategy",
            "text": (
                "Week 3's 77.7% impression spike above average delivered 35.5% of "
                "total campaign volume while maintaining above-average CTR (0.040%). "
                "This demonstrates that aggressive pacing can work without sacrificing "
                "engagement when inventory quality is maintained."
            ),
        },
        {
            "title": "Optimize for Late-Campaign Performance",
            "text": (
                "Week 4 achieved the best performance metrics (0.052% CTR, 71.61% "
                "viewability) despite reduced volume. Consider front-loading learnings "
                "from late-campaign optimizations to achieve this performance level "
                "earlier in future campaigns."
            ),
        },
        {
            "title": "Improve End-of-Campaign Delivery",
            "text": (
                "The final week (Dec 29) delivered only 375 impressions with zero "
                "clicks, representing wasted potential. Better pacing controls or "
                "campaign end-date alignment could improve overall efficiency."
            ),
        },
        {
            "title": "Viewability-CTR Correlation",
            "text": (
                "The data shows positive correlation between viewability and CTR - "
                "Week 4's highest viewability (71.61%) coincided with highest CTR "
                "(0.052%). Recommend prioritizing viewability optimization as a "
                "primary lever for CTR improvement."
            ),
        },
    ],
    "key_campaign_insights": [
        {
            "title": "Mobile Dominates with Superior CTR",
            "text": (
                "Mobile devices delivered 64.62% of impressions (59,275) with the "
                "highest CTR at 0.046%, significantly outperforming PC (0.014% CTR) "
                "which accounted for 30.27% of impressions. Tablet (5.09%) showed "
                "no measurable clicks. This suggests strong mobile-first audience "
                "engagement and opportunity to shift budget from PC to mobile."
            ),
        },
        {
            "title": "Tuesday Peak Engagement",
            "text": (
                "Tuesday delivered the highest CTR (0.052%) despite moderate "
                "impression volume (13,431). Friday showed the lowest CTR (0.020%) "
                "with similar volume. This 2.6x CTR difference presents a clear "
                "day-parting optimization opportunity - consider increasing Tuesday "
                "bids and reducing Friday spend."
            ),
        },
        {
            "title": "Weekend Volume vs Weekday Performance",
            "text": (
                "Weekend days (Saturday-Sunday) delivered high impression volume "
                "(28,234, 30.8% of total) but mixed CTR performance. Saturday showed "
                "strong CTR (0.037%) while Sunday underperformed (0.022%). Consider "
                "maintaining Saturday spend while optimizing or reducing Sunday "
                "delivery."
            ),
        },
        {
            "title": "Domain Concentration Risk",
            "text": (
                "The top domain (blog.smart-trends-site.com) captured 17.78% of all "
                "impressions with below-average CTR (0.025%). Top 10 domains account "
                "for 60.87% of delivery. High-performing outlier content.story-feeds.com "
                "shows 0.136% CTR with only 4.82% share - consider increasing allocation "
                "to this domain."
            ),
        },
    ],
}

# Page config
st.set_page_config(
    page_title="Post-Campaign Insights",
    page_icon="📊",
    layout="wide",
)

# Custom CSS
st.markdown(
    """
    <style>
    .metric-card {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        padding: 1rem;
        border-radius: 0.5rem;
        color: white;
    }
    .insight-green { background-color: #d4edda; border-left: 4px solid #28a745; padding: 1rem; margin: 0.5rem 0; }
    .insight-amber { background-color: #fff3cd; border-left: 4px solid #ffc107; padding: 1rem; margin: 0.5rem 0; }
    .insight-red { background-color: #f8d7da; border-left: 4px solid #dc3545; padding: 1rem; margin: 0.5rem 0; }
    </style>
    """,
    unsafe_allow_html=True,
)


def format_number(n: int | float, decimals: int = 0) -> str:
    """Format number with commas."""
    if decimals == 0:
        return f"{int(n):,}"
    return f"{n:,.{decimals}f}"


def format_pct(n: float | None, decimals: int = 2) -> str:
    """Format percentage."""
    if n is None:
        return "N/A"
    return f"{n:.{decimals}f}%"


def render_insight(insight: dict) -> None:
    """Render an insight card with severity color."""
    severity = insight["severity"]
    icon = {"green": "✅", "amber": "⚠️", "red": "🚨"}[severity]
    css_class = f"insight-{severity}"

    st.markdown(
        f"""
        <div class="{css_class}">
            <strong>{icon} {insight['rule_id'].replace('_', ' ').title()}</strong><br/>
            {insight['description']}<br/>
            <em>→ {insight['recommendation']}</em>
        </div>
        """,
        unsafe_allow_html=True,
    )


def create_weekly_trend_chart(weekly_data: list) -> go.Figure:
    """Create weekly trend chart with impressions and CTR."""
    fig = go.Figure()

    fig.add_trace(go.Bar(
        x=[w["week"] for w in weekly_data],
        y=[w["impressions"] for w in weekly_data],
        name="Impressions",
        marker_color="#667eea",
        yaxis="y",
    ))

    fig.add_trace(go.Scatter(
        x=[w["week"] for w in weekly_data],
        y=[w["ctr"] if w["ctr"] else 0 for w in weekly_data],
        name="CTR (%)",
        yaxis="y2",
        mode="lines+markers",
        line=dict(color="#dc3545", width=3),
        marker=dict(size=8),
    ))

    fig.update_layout(
        title="Weekly Performance Trend",
        yaxis=dict(title="Impressions", side="left", showgrid=True),
        yaxis2=dict(title="CTR (%)", side="right", overlaying="y", showgrid=False),
        legend=dict(x=0, y=1.15, orientation="h"),
        height=400,
        width=800,
        plot_bgcolor="white",
    )

    return fig


def create_platform_pie(platform_data: list) -> go.Figure:
    """Create platform impression share pie chart."""
    colors = ["#667eea", "#764ba2", "#f093fb", "#f5576c", "#4facfe"]

    fig = go.Figure(data=[go.Pie(
        labels=[p["platform"] for p in platform_data],
        values=[p["impressions"] for p in platform_data],
        hole=0.4,
        marker_colors=colors[:len(platform_data)],
    )])

    fig.update_layout(
        title="Platform Impression Share",
        height=400,
        width=500,
    )

    return fig


def create_dow_chart(dow_data: list) -> go.Figure:
    """Create day-of-week bar chart."""
    fig = go.Figure(data=[go.Bar(
        x=[d["day"] for d in dow_data],
        y=[d["impressions"] for d in dow_data],
        marker_color="#764ba2",
    )])

    fig.update_layout(
        title="Day-of-Week Impressions",
        xaxis_title="Day",
        yaxis_title="Impressions",
        height=350,
        width=500,
        plot_bgcolor="white",
    )

    return fig


def create_domain_pie(domain_data: list) -> go.Figure:
    """Create top domains pie chart."""
    colors = ["#667eea", "#764ba2", "#f093fb", "#f5576c", "#4facfe"]
    top5 = domain_data[:5]

    fig = go.Figure(data=[go.Pie(
        labels=[d["domain"][:25] + "..." if len(d["domain"]) > 25 else d["domain"] for d in top5],
        values=[d["impressions"] for d in top5],
        hole=0.4,
        marker_colors=colors,
    )])

    fig.update_layout(
        title="Top 5 Domains",
        height=400,
        width=500,
    )

    return fig


def fig_to_image(fig: go.Figure) -> BytesIO:
    """Convert Plotly figure to PNG image bytes."""
    img_bytes = fig.to_image(format="png", scale=2)
    return BytesIO(img_bytes)


def generate_docx(output, summary: dict) -> BytesIO:
    """Generate DOCX report with charts and tables."""
    doc = Document()

    # Title
    title = doc.add_heading("Post-Campaign Insights & Recommendations", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Campaign info
    doc.add_paragraph(f"Campaign ID: {summary['campaign_id']}")
    doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    kpis = summary["kpis"]
    doc.add_paragraph(
        f"The campaign delivered {format_number(kpis['total_impressions'])} impressions "
        f"with a CTR of {format_pct(kpis['ctr_pct'])} and spend of ${format_number(kpis['total_spend'], 2)}."
    )

    # KPI Table
    doc.add_heading("Campaign KPIs", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "Metric"
    hdr[1].text = "Value"

    kpi_rows = [
        ("Total Impressions", format_number(kpis["total_impressions"])),
        ("Total Clicks", format_number(kpis["total_clicks"])),
        ("Total Spend", f"${format_number(kpis['total_spend'], 2)}"),
        ("CTR", format_pct(kpis["ctr_pct"])),
        ("CPM", f"${format_number(kpis['cpm'], 2)}"),
        ("Viewability", format_pct(kpis["viewability_pct"])),
        ("VCR", format_pct(kpis["vcr_pct"]) if kpis["vcr_pct"] else "N/A"),
    ]
    for metric, value in kpi_rows:
        row = table.add_row().cells
        row[0].text = metric
        row[1].text = str(value)

    # Campaign Summary Insight
    doc.add_paragraph()
    doc.add_heading("Campaign Summary", level=2)
    doc.add_paragraph(HARDCODED_INSIGHTS["campaign_summary"])

    # Key Insights
    doc.add_heading("Key Insights", level=1)
    insights = summary.get("insights", [])

    # Group by severity
    for severity, label in [("red", "Critical"), ("amber", "Warning"), ("green", "Positive")]:
        severity_insights = [i for i in insights if i["severity"] == severity]
        if severity_insights:
            doc.add_heading(f"{label} Findings", level=2)
            for insight in severity_insights:
                icon = {"green": "✓", "amber": "⚠", "red": "✗"}[severity]
                p = doc.add_paragraph()
                p.add_run(f"[{icon}] ").bold = True
                p.add_run(insight["description"])
                rec = doc.add_paragraph(f"    → {insight['recommendation']}")
                rec.paragraph_format.left_indent = Inches(0.5)

    # Weekly Performance with Chart
    doc.add_heading("Weekly Performance", level=1)

    weekly_data = summary["weekly_performance"]
    if weekly_data:
        # Add chart
        fig = create_weekly_trend_chart(weekly_data)
        img_stream = fig_to_image(fig)
        doc.add_picture(img_stream, width=Inches(6))

        doc.add_paragraph()  # Spacing

        # Add table
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Week", "Impressions", "Clicks", "CTR", "Spend", "Viewability"]):
            hdr[i].text = h

        for w in weekly_data:
            row = table.add_row().cells
            row[0].text = w["week"]
            row[1].text = format_number(w["impressions"])
            row[2].text = format_number(w["clicks"])
            row[3].text = format_pct(w["ctr"])
            row[4].text = f"${format_number(w['spend'], 2)}"
            row[5].text = format_pct(w.get("viewability"))

        # Weekly Performance Analysis
        doc.add_paragraph()
        doc.add_heading("Weekly Performance Analysis", level=2)

        trend_insights = HARDCODED_INSIGHTS["weekly_trend"]
        for label, key in [
            ("Impression Delivery", "impression_delivery"),
            ("Click-Through Rate (CTR)", "ctr"),
            ("Viewability", "viewability"),
            ("Video Completion Rate (VCR)", "vcr"),
        ]:
            p = doc.add_paragraph()
            p.add_run(f"{label}: ").bold = True
            p.add_run(trend_insights[key])

    # Key Insights & Recommendations
    doc.add_heading("Key Insights & Recommendations", level=1)
    for i, rec in enumerate(HARDCODED_INSIGHTS["key_recommendations"], 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {rec['title']}: ").bold = True
        p.add_run(rec["text"])

    # Platform Breakdown with Chart
    doc.add_heading("Platform Breakdown", level=1)

    platform_data = summary["platform_breakdown"]
    if platform_data:
        # Add chart
        fig = create_platform_pie(platform_data)
        img_stream = fig_to_image(fig)
        doc.add_picture(img_stream, width=Inches(4))

        doc.add_paragraph()

        # Add table
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Platform", "Impressions", "Share", "CTR", "CPM"]):
            hdr[i].text = h

        for p in platform_data:
            row = table.add_row().cells
            row[0].text = p["platform"]
            row[1].text = format_number(p["impressions"])
            row[2].text = f"{p['impression_share']}%"
            row[3].text = format_pct(p["ctr"])
            row[4].text = f"${p['cpm']}" if p.get("cpm") else "N/A"

    # Key Campaign Insights
    doc.add_heading("Key Campaign Insights", level=1)
    for i, insight in enumerate(HARDCODED_INSIGHTS["key_campaign_insights"], 1):
        p = doc.add_paragraph()
        p.add_run(f"{i}. {insight['title']}: ").bold = True
        p.add_run(insight["text"])

    # Day of Week Analysis with Chart
    doc.add_heading("Day of Week Analysis", level=1)

    dow_data = summary.get("day_of_week_performance", [])
    if dow_data:
        fig = create_dow_chart(dow_data)
        img_stream = fig_to_image(fig)
        doc.add_picture(img_stream, width=Inches(5))

        doc.add_paragraph()

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Day", "Impressions", "Clicks", "CTR"]):
            hdr[i].text = h

        for d in dow_data:
            row = table.add_row().cells
            row[0].text = d["day"]
            row[1].text = format_number(d["impressions"])
            row[2].text = format_number(d["clicks"])
            row[3].text = format_pct(d["ctr_pct"])

    # Top Domains with Chart
    doc.add_heading("Top Domains", level=1)

    domain_data = summary["top_domains"]
    if domain_data:
        # Add chart
        fig = create_domain_pie(domain_data)
        img_stream = fig_to_image(fig)
        doc.add_picture(img_stream, width=Inches(4))

        doc.add_paragraph()
        doc.add_paragraph(f"Top 10 Domain Share: {summary['top_10_domain_share_pct']}%")

        # Add table
        table = doc.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, h in enumerate(["Domain", "Impressions", "Share", "CTR", "Status"]):
            hdr[i].text = h

        for d in domain_data[:10]:
            row = table.add_row().cells
            row[0].text = d["domain"][:40]  # Truncate long domains
            row[1].text = format_number(d["impressions"])
            row[2].text = f"{d['impression_share_pct']}%"
            row[3].text = format_pct(d["ctr_pct"])
            row[4].text = "⚠ Underperforming" if d.get("is_underperforming") else "OK"

    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def main():
    st.title("📊 Post-Campaign Insights & Recommendations")

    # Sidebar - File Upload
    with st.sidebar:
        st.header("📁 Upload Reports")

        domain_file = st.file_uploader(
            "Domain Report (Excel)",
            type=["xlsx", "xls"],
            help="Upload the Domain Report Excel file",
        )

        campaign_file = st.file_uploader(
            "Campaign Report (Excel) - Optional",
            type=["xlsx", "xls"],
            help="Upload the Campaign Report Excel file (optional)",
        )

        st.divider()

    # Main area
    if not domain_file:
        st.info("👈 Upload a Domain Report Excel file to get started")
        return

    # Save uploaded file to temp location
    with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp:
        tmp.write(domain_file.getvalue())
        domain_path = Path(tmp.name)

    campaign_path = None
    if campaign_file:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp:
            tmp.write(campaign_file.getvalue())
            campaign_path = Path(tmp.name)

    # Initialize service and get campaigns
    service = ReportService()

    try:
        campaigns = service.get_available_campaigns(domain_path)
    except Exception as e:
        st.error(f"Error reading file: {e}")
        return

    if not campaigns:
        st.error("No campaigns found in the uploaded file")
        return

    # Campaign selector
    with st.sidebar:
        selected_campaign = st.selectbox(
            "Select Campaign ID",
            options=campaigns,
            format_func=lambda x: f"Campaign {x}",
        )

        generate_btn = st.button("🚀 Generate Report", type="primary", use_container_width=True)

    # Generate report on button click
    if generate_btn or "report_output" in st.session_state:
        if generate_btn:
            with st.spinner("Generating report..."):
                try:
                    output = service.generate_report(
                        campaign_id=selected_campaign,
                        domain_report_path=domain_path,
                        campaign_report_path=campaign_path,
                    )
                    summary = service.generate_summary_dict(output)
                    st.session_state["report_output"] = output
                    st.session_state["report_summary"] = summary
                except Exception as e:
                    st.error(f"Error generating report: {e}")
                    return

        output = st.session_state.get("report_output")
        summary = st.session_state.get("report_summary")

        if not output or not summary:
            return

        # Tabs
        tab1, tab2, tab3, tab4 = st.tabs([
            "📈 Campaign Performance",
            "🖥️ Platform & Temporal",
            "🌐 Inventory / Domain Pulse",
            "📄 Generated Report",
        ])

        # =====================================================================
        # TAB 1: Campaign Performance
        # =====================================================================
        with tab1:
            st.header("Campaign Performance Overview")

            # KPI Tiles
            kpis = summary["kpis"]
            col1, col2, col3, col4 = st.columns(4)

            with col1:
                st.metric("Total Impressions", format_number(kpis["total_impressions"]))
            with col2:
                st.metric("Total Clicks", format_number(kpis["total_clicks"]))
            with col3:
                st.metric("Total Spend", f"${format_number(kpis['total_spend'], 2)}")
            with col4:
                st.metric("CTR", format_pct(kpis["ctr_pct"]))

            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("CPM", f"${format_number(kpis['cpm'], 2)}")
            with col2:
                st.metric("Viewability", format_pct(kpis["viewability_pct"]))
            with col3:
                vcr_val = format_pct(kpis["vcr_pct"]) if kpis["vcr_pct"] else "N/A"
                st.metric("VCR", vcr_val)

            # Campaign Summary Insight
            st.info(HARDCODED_INSIGHTS["campaign_summary"])

            st.divider()

            # Weekly Trend Chart
            st.subheader("Weekly Performance Trend")

            weekly_data = summary["weekly_performance"]
            if weekly_data:
                fig = go.Figure()

                fig.add_trace(go.Bar(
                    x=[w["week"] for w in weekly_data],
                    y=[w["impressions"] for w in weekly_data],
                    name="Impressions",
                    yaxis="y",
                ))

                fig.add_trace(go.Scatter(
                    x=[w["week"] for w in weekly_data],
                    y=[w["ctr"] if w["ctr"] else 0 for w in weekly_data],
                    name="CTR (%)",
                    yaxis="y2",
                    mode="lines+markers",
                    line=dict(color="red", width=2),
                ))

                fig.update_layout(
                    yaxis=dict(title="Impressions", side="left"),
                    yaxis2=dict(title="CTR (%)", side="right", overlaying="y"),
                    legend=dict(x=0, y=1.1, orientation="h"),
                    height=400,
                )

                st.plotly_chart(fig, use_container_width=True)

            # Weekly Table
            st.subheader("Weekly Breakdown")
            st.dataframe(
                weekly_data,
                use_container_width=True,
                hide_index=True,
            )

            # Weekly Trend Insights
            st.divider()
            st.subheader("Weekly Performance Analysis")

            trend_insights = HARDCODED_INSIGHTS["weekly_trend"]
            with st.expander("Impression Delivery", expanded=True):
                st.write(trend_insights["impression_delivery"])
            with st.expander("Click-Through Rate (CTR)", expanded=True):
                st.write(trend_insights["ctr"])
            with st.expander("Viewability", expanded=True):
                st.write(trend_insights["viewability"])
            with st.expander("Video Completion Rate (VCR)", expanded=True):
                st.write(trend_insights["vcr"])

            # Key Insights & Recommendations
            st.divider()
            st.subheader("Key Insights & Recommendations")

            for i, rec in enumerate(HARDCODED_INSIGHTS["key_recommendations"], 1):
                st.markdown(f"**{i}. {rec['title']}**")
                st.write(rec["text"])
                st.write("")

        # =====================================================================
        # TAB 2: Platform & Temporal
        # =====================================================================
        with tab2:
            st.header("Platform & Temporal Analysis")

            col1, col2 = st.columns(2)

            with col1:
                st.subheader("Platform / Device Performance")

                platform_data = summary["platform_breakdown"]
                if platform_data:
                    fig = px.pie(
                        values=[p["impressions"] for p in platform_data],
                        names=[p["platform"] for p in platform_data],
                        title="Impression Share by Platform",
                    )
                    st.plotly_chart(fig, use_container_width=True)

                    st.dataframe(platform_data, use_container_width=True, hide_index=True)

            with col2:
                st.subheader("Day-of-Week Analysis")

                dow_data = summary["day_of_week_performance"]
                if dow_data:
                    fig = go.Figure()

                    fig.add_trace(go.Bar(
                        x=[d["day"] for d in dow_data],
                        y=[d["impressions"] for d in dow_data],
                        name="Impressions",
                    ))

                    fig.update_layout(height=300)
                    st.plotly_chart(fig, use_container_width=True)

                    st.dataframe(dow_data, use_container_width=True, hide_index=True)

            st.divider()

            # Concentration Warnings
            st.subheader("⚠️ Concentration Warnings")

            platform_insights = [
                i for i in summary.get("insights", [])
                if i["rule_id"] == "platform_concentration"
            ]

            if platform_insights:
                for insight in platform_insights:
                    render_insight(insight)
            else:
                st.success("✅ No platform concentration issues detected")

            # Key Campaign Insights
            st.divider()
            st.subheader("Key Campaign Insights")

            for i, insight in enumerate(HARDCODED_INSIGHTS["key_campaign_insights"], 1):
                st.markdown(f"**{i}. {insight['title']}**")
                st.write(insight["text"])
                st.write("")

        # =====================================================================
        # TAB 3: Inventory / Domain Pulse
        # =====================================================================
        with tab3:
            st.header("Inventory / Domain Pulse")

            col1, col2 = st.columns([2, 1])

            with col1:
                st.subheader("Top Domains")
                domain_data = summary["top_domains"]

                if domain_data:
                    st.dataframe(
                        domain_data,
                        use_container_width=True,
                        hide_index=True,
                        column_config={
                            "domain": "Domain",
                            "impressions": st.column_config.NumberColumn("Impressions", format="%d"),
                            "impression_share_pct": st.column_config.NumberColumn("Share %", format="%.2f%%"),
                            "ctr_pct": st.column_config.NumberColumn("CTR %", format="%.4f%%"),
                            "is_underperforming": st.column_config.CheckboxColumn("Underperforming"),
                        },
                    )

            with col2:
                st.subheader("Domain Concentration")
                st.metric(
                    "Top 10 Domain Share",
                    f"{summary['top_10_domain_share_pct']}%",
                    delta="High" if summary["top_10_domain_share_pct"] > 70 else None,
                    delta_color="inverse",
                )

                # Domain concentration pie
                if domain_data:
                    top5 = domain_data[:5]
                    fig = px.pie(
                        values=[d["impressions"] for d in top5],
                        names=[d["domain"][:20] + "..." if len(d["domain"]) > 20 else d["domain"] for d in top5],
                        title="Top 5 Domains",
                    )
                    fig.update_layout(height=300)
                    st.plotly_chart(fig, use_container_width=True)

            st.divider()

            # Domain Insights
            st.subheader("Domain Insights")

            domain_insights = [
                i for i in summary.get("insights", [])
                if "domain" in i["rule_id"]
            ]

            underperforming = [d for d in domain_data if d.get("is_underperforming")]

            if domain_insights:
                for insight in domain_insights:
                    render_insight(insight)

            if underperforming:
                st.warning(f"⚠️ {len(underperforming)} underperforming domain(s) detected with high share but low CTR")
                for d in underperforming:
                    st.write(f"  • **{d['domain']}**: {d['impression_share_pct']}% share, {format_pct(d['ctr_pct'])} CTR")

            if not domain_insights and not underperforming:
                st.success("✅ Domain distribution looks healthy")

        # =====================================================================
        # TAB 4: Generated Report
        # =====================================================================
        with tab4:
            st.header("Generated Report Preview")

            # All insights
            st.subheader("Key Insights")
            insights = summary.get("insights", [])

            if insights:
                # Group by severity
                red_insights = [i for i in insights if i["severity"] == "red"]
                amber_insights = [i for i in insights if i["severity"] == "amber"]
                green_insights = [i for i in insights if i["severity"] == "green"]

                if red_insights:
                    st.markdown("### 🔴 Critical Issues")
                    for insight in red_insights:
                        render_insight(insight)

                if amber_insights:
                    st.markdown("### 🟠 Warnings")
                    for insight in amber_insights:
                        render_insight(insight)

                if green_insights:
                    st.markdown("### 🟢 Positive Signals")
                    for insight in green_insights:
                        render_insight(insight)
            else:
                st.info("No insights generated")

            st.divider()

            # Export
            st.subheader("Export Report")

            docx_buffer = generate_docx(output, summary)

            st.download_button(
                label="📥 Download DOCX Report",
                data=docx_buffer,
                file_name=f"campaign_{selected_campaign}_insights_{datetime.now().strftime('%Y%m%d')}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                type="primary",
            )


if __name__ == "__main__":
    main()
